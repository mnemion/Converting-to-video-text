from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Body
from fastapi.middleware.cors import CORSMiddleware
from starlette.staticfiles import StaticFiles
from fastapi.responses import StreamingResponse, FileResponse
import os
import shutil
import uuid
from io import BytesIO
from dotenv import load_dotenv
import csv
import json
import re

# 내부 모듈
from utils.validator import allowed_file, validate_file_size
from tasks.video_processing import extract_audio, convert_wav_to_mp3
from tasks.transcription import TranscriptionService
from celery_app import celery_app
from tasks.async_transcription import transcribe_video_async, transcribe_url_async, download_url_async
from tasks.url_download import download_media_via_ytdlp


load_dotenv()
# .env.local도 있으면 덮어쓰기 로드
try:
    load_dotenv(os.path.join(os.path.dirname(__file__), ".env.local"), override=True)
except Exception:
    pass

app = FastAPI(
    title="Video Transcription API",
    description="동영상 업로드, 전사, 자막 생성 API",
    version="1.0.0",
)

ALLOWED_LANGUAGE_CODES = {"auto", "ko", "en", "ja", "zh", "es", "fr"}

cors_origins = os.getenv("CORS_ORIGINS", "*")
allow_origins = [o.strip() for o in cors_origins.split(",") if o.strip()]
app.add_middleware(
    CORSMiddleware,
    allow_origins=allow_origins or ["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

UPLOAD_FOLDER = "uploads"
OUTPUT_FOLDER = "outputs"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

app.mount("/uploads", StaticFiles(directory=UPLOAD_FOLDER), name="uploads")
app.mount("/outputs", StaticFiles(directory=OUTPUT_FOLDER), name="outputs")


transcription_service = TranscriptionService(model_size=os.getenv("WHISPER_MODEL_SIZE", "base"))


@app.post("/upload")
async def upload_video(file: UploadFile = File(...)):
    if not file or file.filename == "":
        raise HTTPException(status_code=400, detail="파일이 없습니다")
    if not allowed_file(file.filename):
        raise HTTPException(status_code=400, detail="지원하지 않는 파일 형식입니다")
    if not validate_file_size(file.file):
        raise HTTPException(status_code=400, detail="파일이 너무 큽니다 (최대 500MB)")

    filename = os.path.basename(file.filename)
    filepath = os.path.join(UPLOAD_FOLDER, filename)
    with open(filepath, "wb") as out_file:
        shutil.copyfileobj(file.file, out_file)
    return {"message": "업로드 성공", "filename": filename, "filepath": f"uploads/{filename}"}


@app.post("/transcribe")
async def transcribe_video(file: UploadFile = File(...), language: str = Form("ko"), model: str = Form(None), diarize: str = Form(None)):
    if not file or file.filename == "":
        raise HTTPException(status_code=400, detail="파일이 없습니다")
    language_code = (language or "ko").lower()
    if language_code not in ALLOWED_LANGUAGE_CODES:
        raise HTTPException(status_code=400, detail="지원하지 않는 언어 코드입니다 (ko,en,ja,zh,es,fr)")
    # auto는 Whisper 자동 감지(None)로 위임
    effective_lang = None if language_code == "auto" else language_code

    job_id = str(uuid.uuid4())

    video_filename = f"{job_id}_{os.path.basename(file.filename)}"
    video_path = os.path.join(UPLOAD_FOLDER, video_filename)
    with open(video_path, "wb") as f:
        shutil.copyfileobj(file.file, f)

    # 원본 파일명 메타 저장
    try:
        meta_path = os.path.join(OUTPUT_FOLDER, f"{job_id}.json")
        with open(meta_path, "w", encoding="utf-8") as mf:
            json.dump({"job_id": job_id, "original_filename": os.path.basename(file.filename)}, mf, ensure_ascii=False)
    except Exception:
        pass

    audio_path = os.path.join(UPLOAD_FOLDER, f"{job_id}.wav")
    success, result = extract_audio(video_path, audio_path)
    if not success:
        raise HTTPException(status_code=500, detail=result)

    # 동기 전사에서도 MP3 생성 (출력 폴더)
    mp3_out = os.path.join(OUTPUT_FOLDER, f"{job_id}.mp3")
    try:
        os.makedirs(OUTPUT_FOLDER, exist_ok=True)
        ok, mp3_res = convert_wav_to_mp3(audio_path, mp3_out)
        if not ok:
            # 변환 실패는 치명적이지 않으므로 로그 수준으로만 처리
            mp3_out = None
    except Exception:
        mp3_out = None

    # 요청 단위 모델 스위치(선택): 모델명이 다르면 임시 인스턴스 생성
    model_size = (model or os.getenv("WHISPER_MODEL_SIZE", "base")).strip()
    svc = transcription_service if model_size == os.getenv("WHISPER_MODEL_SIZE", "base") else TranscriptionService(model_size=model_size)
    transcription_result = svc.transcribe(audio_path, effective_lang)
    if not transcription_result.get("success"):
        raise HTTPException(status_code=500, detail=transcription_result.get("error", "전사 실패"))

    output_txt = os.path.join(OUTPUT_FOLDER, f"{job_id}.txt")
    output_srt = os.path.join(OUTPUT_FOLDER, f"{job_id}.srt")
    transcription_service.save_transcription(transcription_result, output_txt)
    # 화자 분리(선택)
    do_diarize = str(diarize or "").lower() in ("1", "true", "yes", "on")
    if do_diarize:
        try:
            from tasks.diarization import diarize_audio, write_srt_with_speakers  # type: ignore
            speakers = diarize_audio(audio_path)
            ok = write_srt_with_speakers(transcription_result["segments"], speakers, output_srt)
            if not ok:
                transcription_service.create_srt(transcription_result["segments"], output_srt)
        except Exception:
            transcription_service.create_srt(transcription_result["segments"], output_srt)
    else:
        transcription_service.create_srt(transcription_result["segments"], output_srt)

    try:
        os.remove(video_path)
        os.remove(audio_path)
    except Exception:
        pass

    return {
        "job_id": job_id,
        "text": transcription_result["text"],
        "txt_file": f"outputs/{job_id}.txt",
        "srt_file": f"outputs/{job_id}.srt",
        "language": transcription_result["language"],
        "audio_mp3": f"outputs/{job_id}.mp3" if mp3_out and os.path.exists(mp3_out) else None,
    }


@app.post("/transcribe-async")
async def transcribe_video_async_endpoint(file: UploadFile = File(...), language: str = Form("ko"), model: str = Form(None), diarize: str = Form(None)):
    if not file or file.filename == "":
        raise HTTPException(status_code=400, detail="파일이 없습니다")
    language_code = (language or "ko").lower()
    if language_code not in ALLOWED_LANGUAGE_CODES:
        raise HTTPException(status_code=400, detail="지원하지 않는 언어 코드입니다 (ko,en,ja,zh,es,fr)")
    effective_lang = None if language_code == "auto" else language_code

    job_id = str(uuid.uuid4())
    video_filename = f"{job_id}_{os.path.basename(file.filename)}"
    video_path = os.path.join(UPLOAD_FOLDER, video_filename)
    with open(video_path, "wb") as f:
        shutil.copyfileobj(file.file, f)

    # 원본 파일명 메타 저장 (비동기용)
    try:
        meta_path = os.path.join(OUTPUT_FOLDER, f"{job_id}.json")
        with open(meta_path, "w", encoding="utf-8") as mf:
            json.dump({"job_id": job_id, "original_filename": os.path.basename(file.filename)}, mf, ensure_ascii=False)
    except Exception:
        pass

    # diarize 플래그 전달 (문자열 true/false 수용)
    do_diarize = str(diarize or "").lower() in ("1", "true", "yes", "on")
    # 모델 크기(프론트에서 전달된 값 우선)
    model_size = (model or os.getenv("WHISPER_MODEL_SIZE", "base")).strip()
    task = transcribe_video_async.delay(video_path, effective_lang, do_diarize, model_size)
    return {"job_id": job_id, "task_id": task.id, "status": "processing"}


@app.post("/transcribe-url-async")
async def transcribe_url_async_endpoint(url: str = Form(...), language: str = Form("ko"), model: str = Form(None), diarize: str = Form(None)):
    if not isinstance(url, str) or not re.match(r"^https?://", url.strip()):
        raise HTTPException(status_code=400, detail="유효한 URL이 아닙니다")
    language_code = (language or "ko").lower()
    if language_code not in ALLOWED_LANGUAGE_CODES:
        raise HTTPException(status_code=400, detail="지원하지 않는 언어 코드입니다 (ko,en,ja,zh,es,fr)")
    effective_lang = None if language_code == "auto" else language_code

    job_id = str(uuid.uuid4())
    do_diarize = str(diarize or "").lower() in ("1", "true", "yes", "on")
    model_size = (model or os.getenv("WHISPER_MODEL_SIZE", "base")).strip()

    # 메타에 원본 URL 저장
    try:
        meta_path = os.path.join(OUTPUT_FOLDER, f"{job_id}.json")
        with open(meta_path, "w", encoding="utf-8") as mf:
            json.dump({"job_id": job_id, "source_url": url}, mf, ensure_ascii=False)
    except Exception:
        pass

    task = transcribe_url_async.delay(url, job_id, effective_lang, do_diarize, model_size)
    return {"job_id": job_id, "task_id": task.id, "status": "processing"}


@app.post("/transcribe-url")
async def transcribe_url(url: str = Form(...), language: str = Form("ko"), model: str = Form(None), diarize: str = Form(None)):
    if not isinstance(url, str) or not re.match(r"^https?://", url.strip()):
        raise HTTPException(status_code=400, detail="유효한 URL이 아닙니다")
    language_code = (language or "ko").lower()
    if language_code not in ALLOWED_LANGUAGE_CODES:
        raise HTTPException(status_code=400, detail="지원하지 않는 언어 코드입니다 (ko,en,ja,zh,es,fr)")
    effective_lang = None if language_code == "auto" else language_code

    job_id = str(uuid.uuid4())

    # 다운로드
    ok, dl = download_media_via_ytdlp(url.strip(), job_id, UPLOAD_FOLDER)
    if not ok:
        raise HTTPException(status_code=400, detail=str(dl.get("error", "다운로드 실패")))
    video_path = dl.get("path")
    original_title = str(dl.get("title") or os.path.basename(video_path))

    # 메타 저장
    try:
        meta_path = os.path.join(OUTPUT_FOLDER, f"{job_id}.json")
        with open(meta_path, "w", encoding="utf-8") as mf:
            json.dump({"job_id": job_id, "original_filename": original_title, "source_url": url}, mf, ensure_ascii=False)
    except Exception:
        pass

    # 오디오 추출
    audio_path = os.path.join(UPLOAD_FOLDER, f"{job_id}.wav")
    success, result = extract_audio(video_path, audio_path)
    if not success:
        raise HTTPException(status_code=500, detail=result)

    # MP3 생성 (출력 폴더)
    mp3_out = os.path.join(OUTPUT_FOLDER, f"{job_id}.mp3")
    try:
        os.makedirs(OUTPUT_FOLDER, exist_ok=True)
        ok2, _ = convert_wav_to_mp3(audio_path, mp3_out)
        if not ok2:
            mp3_out = None
    except Exception:
        mp3_out = None

    # 모델 선택 및 전사
    model_size = (model or os.getenv("WHISPER_MODEL_SIZE", "base")).strip()
    svc = transcription_service if model_size == os.getenv("WHISPER_MODEL_SIZE", "base") else TranscriptionService(model_size=model_size)
    transcription_result = svc.transcribe(audio_path, effective_lang)
    if not transcription_result.get("success"):
        raise HTTPException(status_code=500, detail=transcription_result.get("error", "전사 실패"))

    output_txt = os.path.join(OUTPUT_FOLDER, f"{job_id}.txt")
    output_srt = os.path.join(OUTPUT_FOLDER, f"{job_id}.srt")
    transcription_service.save_transcription(transcription_result, output_txt)
    transcription_service.create_srt(transcription_result["segments"], output_srt)

    try:
        os.remove(video_path)
        os.remove(audio_path)
    except Exception:
        pass

    return {
        "job_id": job_id,
        "text": transcription_result["text"],
        "txt_file": f"outputs/{job_id}.txt",
        "srt_file": f"outputs/{job_id}.srt",
        "language": transcription_result["language"],
        "audio_mp3": f"outputs/{job_id}.mp3" if mp3_out and os.path.exists(mp3_out) else None,
        "original_filename": original_title,
        "source_url": url,
    }


@app.post("/fetch-url-async")
async def fetch_url_async(url: str = Form(...)):
    if not isinstance(url, str) or not re.match(r"^https?://", url.strip()):
        raise HTTPException(status_code=400, detail="유효한 URL이 아닙니다")
    job_id = str(uuid.uuid4())
    # 초기 메타
    try:
        with open(os.path.join(OUTPUT_FOLDER, f"{job_id}.json"), "w", encoding="utf-8") as mf:
            json.dump({"job_id": job_id, "source_url": url}, mf, ensure_ascii=False)
    except Exception:
        pass
    task = download_url_async.delay(url, job_id)
    return {"job_id": job_id, "task_id": task.id, "status": "processing"}


@app.post("/transcribe-downloaded-async")
async def transcribe_downloaded_async(job_id: str = Form(...), language: str = Form("ko"), model: str = Form(None), diarize: str = Form(None)):
    if not job_id or not isinstance(job_id, str):
        raise HTTPException(status_code=400, detail="유효한 job_id")
    language_code = (language or "ko").lower()
    if language_code not in ALLOWED_LANGUAGE_CODES:
        raise HTTPException(status_code=400, detail="지원하지 않는 언어 코드입니다 (ko,en,ja,zh,es,fr)")
    effective_lang = None if language_code == "auto" else language_code
    do_diarize = str(diarize or "").lower() in ("1", "true", "yes", "on")
    model_size = (model or os.getenv("WHISPER_MODEL_SIZE", "base")).strip()

    # 업로드 폴더에서 해당 job_id로 시작하는 파일 검색
    try:
        candidates = [
            os.path.join(UPLOAD_FOLDER, f) for f in os.listdir(UPLOAD_FOLDER)
            if f.startswith(f"{job_id}_")
        ]
    except Exception:
        candidates = []
    if not candidates:
        raise HTTPException(status_code=404, detail="다운로드된 미디어를 찾을 수 없습니다")
    candidates.sort(key=lambda p: os.path.getmtime(p), reverse=True)
    video_path = candidates[0]

    task = transcribe_video_async.delay(video_path, effective_lang, do_diarize, model_size)
    return {"job_id": job_id, "task_id": task.id, "status": "processing"}

@app.get("/status/{task_id}")
def get_task_status(task_id: str):
    task = celery_app.AsyncResult(task_id)
    if task.state == "PENDING":
        return {"state": task.state, "progress": 0}
    if task.state == "PROGRESS":
        info = task.info or {}
        return {"state": task.state, "progress": info.get("progress", 0)}
    if task.state == "SUCCESS":
        return {"state": task.state, "result": task.result}
    return {"state": task.state, "error": str(task.info)}


# 전사 결과 삭제 (txt/srt 파일 제거)
@app.delete("/transcription/{job_id}")
def delete_transcription(job_id: str):
    targets = [
        os.path.join(OUTPUT_FOLDER, f"{job_id}.txt"),
        os.path.join(OUTPUT_FOLDER, f"{job_id}.srt"),
    ]
    deleted = []
    for path in targets:
        try:
            if os.path.exists(path):
                os.remove(path)
                deleted.append(os.path.basename(path))
        except Exception:
            # 한 파일 실패해도 나머지는 시도
            pass
    return {"job_id": job_id, "deleted": deleted}


# 전사 결과 텍스트 수정
@app.put("/transcription/{job_id}/text")
def update_transcript_text(job_id: str, text: str = Body(..., embed=True)):
    if not isinstance(text, str):
        raise HTTPException(status_code=400, detail="잘못된 본문")
    txt_path = os.path.join(OUTPUT_FOLDER, f"{job_id}.txt")
    try:
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(text)
        return {"ok": True}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/export/audio/{job_id}")
def export_audio(job_id: str):
    """MP3 전용 다운로드. outputs/{job_id}.mp3만 허용. 파일명은 원본 이름 기반."""
    mp3_path = os.path.join(OUTPUT_FOLDER, f"{job_id}.mp3")
    if os.path.exists(mp3_path):
        download_name = f"{job_id}.mp3"
        # 메타에서 원본 파일명 읽어서 mp3 확장자로 교체
        try:
            meta_path = os.path.join(OUTPUT_FOLDER, f"{job_id}.json")
            if os.path.exists(meta_path):
                with open(meta_path, "r", encoding="utf-8") as mf:
                    meta = json.load(mf)
                orig = os.path.basename(str(meta.get("original_filename", "")))
                base = os.path.splitext(orig)[0].strip() or job_id
                download_name = base + ".mp3"
        except Exception:
            download_name = f"{job_id}.mp3"
        return FileResponse(mp3_path, media_type="audio/mpeg", filename=download_name)
    raise HTTPException(status_code=404, detail="오디오 파일이 없습니다")


# SRT 파서 (간단, 화자 라벨 인식)
def parse_srt_entries(job_id: str):
    srt_path = os.path.join(OUTPUT_FOLDER, f"{job_id}.srt")
    if not os.path.exists(srt_path):
        return []
    with open(srt_path, "r", encoding="utf-8") as f:
        data = f.read()
    blocks = re.split(r"\n\s*\n", data.strip())
    entries = []
    for blk in blocks:
        lines = [ln for ln in blk.splitlines() if ln.strip()]
        if len(lines) < 2:
            continue
        # 첫 줄이 번호일 수 있음
        if re.match(r"^\d+$", lines[0].strip()):
            lines = lines[1:]
        time_line = lines[0]
        m = re.match(r"(\d{2}:\d{2}:\d{2}),\d{3}\s*-->\s*(\d{2}:\d{2}:\d{2}),\d{3}", time_line)
        if not m:
            # 밀리초 포함 매치
            m = re.match(r"(\d{2}:\d{2}:\d{2}),\d{1,3}\s*-->\s*(\d{2}:\d{2}:\d{2}),\d{1,3}", time_line)
        if not m:
            continue
        start_ts, end_ts = m.group(1), m.group(2)
        text = " ".join(lines[1:]).strip()
        # 화자 라벨 추출: [화자 N] / [SPEAKER_1] / '화자 N:' / 'SPEAKER 1:' 지원
        speaker = None
        msp = re.match(r"^\[(?:\s*화자\s*(\d+)|\s*SPEAKER[_\s]*(\d+))\]\s*(.*)$", text, flags=re.IGNORECASE)
        if msp:
            speaker = msp.group(1) or msp.group(2)
            text = msp.group(3).strip()
        else:
            msp2 = re.match(r"^(?:\s*화자\s*(\d+)|\s*SPEAKER[_\s]*(\d+))\s*:\s*(.*)$", text, flags=re.IGNORECASE)
            if msp2:
                speaker = msp2.group(1) or msp2.group(2)
                text = msp2.group(3).strip()
        ent = {"start": start_ts, "end": end_ts, "text": text}
        if speaker:
            try:
                ent["speaker"] = int(speaker)
            except Exception:
                pass
        entries.append(ent)
    return entries

# 내보내기: DOCX (타임스탬프 옵션)
@app.get("/export/docx/{job_id}")
def export_docx(job_id: str, ts: int = 0, spk: int = 0):
    txt_path = os.path.join(OUTPUT_FOLDER, f"{job_id}.txt")
    if not os.path.exists(txt_path):
        raise HTTPException(status_code=404, detail="TXT 파일을 찾을 수 없습니다")
    try:
        from docx import Document  # type: ignore
    except Exception:
        raise HTTPException(status_code=500, detail="DOCX 모듈이 설치되어 있지 않습니다")

    with open(txt_path, "r", encoding="utf-8") as f:
        plain = f.read()

    doc = Document()
    if ts or spk:
        entries = parse_srt_entries(job_id)
        if entries:
            for e in entries:
                parts = []
                if ts and e.get("start"):
                    parts.append(f"[{e['start']}]")
                if spk and e.get("speaker"):
                    parts.append(f"[화자 {e['speaker']}]")
                prefix = (" ".join(parts) + " ") if parts else ""
                doc.add_paragraph(prefix + e["text"])
        else:
            for line in plain.splitlines():
                doc.add_paragraph(line)
    else:
        for line in plain.splitlines():
            doc.add_paragraph(line)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    headers = {"Content-Disposition": f"attachment; filename=transcript_{job_id}.docx"}
    return StreamingResponse(buffer, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers=headers)


# 내보내기: PDF (타임스탬프 옵션)
@app.get("/export/pdf/{job_id}")
def export_pdf(job_id: str, ts: int = 0, spk: int = 0):
    txt_path = os.path.join(OUTPUT_FOLDER, f"{job_id}.txt")
    if not os.path.exists(txt_path):
        raise HTTPException(status_code=404, detail="TXT 파일을 찾을 수 없습니다")
    try:
        from reportlab.lib.pagesizes import A4  # type: ignore
        from reportlab.lib.units import mm  # type: ignore
        from reportlab.pdfbase import pdfmetrics  # type: ignore
        from reportlab.pdfbase.ttfonts import TTFont  # type: ignore
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer  # type: ignore
        from reportlab.lib.styles import ParagraphStyle  # type: ignore
    except Exception:
        raise HTTPException(status_code=500, detail="PDF 모듈이 설치되어 있지 않습니다")

    import html as htmlmod

    with open(txt_path, "r", encoding="utf-8") as f:
        plain = f.read()

    # 폰트 선택 (기존 로직 유지)
    def resolve_font() -> tuple[str | None, int | None]:
        env_path = os.getenv("PDF_FONT_PATH")
        env_index = os.getenv("PDF_FONT_INDEX")
        try:
            env_index_int = int(env_index) if env_index is not None else None
        except Exception:
            env_index_int = None
        if env_path and os.path.exists(env_path):
            return env_path, env_index_int
        candidates: list[tuple[str, int | None]] = [
            ("/usr/share/fonts/opentype/noto/NotoSerifCJKkr-Regular.otf", None),
            ("/usr/share/fonts/opentype/noto/NotoSansCJKkr-Regular.otf", None),
            ("/usr/share/fonts/opentype/noto/NotoSerifCJKsc-Regular.otf", None),
            ("/usr/share/fonts/opentype/noto/NotoSansCJKsc-Regular.otf", None),
            ("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc", 0),
            (r"C:\\Windows\\Fonts\\malgun.ttf", None),
            (r"C:\\Windows\\Fonts\\msyh.ttc", 0),
            (r"C:\\Windows\\Fonts\\meiryo.ttc", 0),
            ("/Library/Fonts/Apple SD Gothic Neo.ttf", None),
            ("/System/Library/Fonts/AppleSDGothicNeo.ttf", None),
        ]
        for path, idx in candidates:
            if os.path.exists(path):
                return path, idx
        return None, None

    def resolve_fallback_font() -> tuple[str | None, int | None]:
        f_path = os.getenv("PDF_FONT_FALLBACK_PATH")
        f_idx = os.getenv("PDF_FONT_FALLBACK_INDEX")
        try:
            f_idx_int = int(f_idx) if f_idx is not None else None
        except Exception:
            f_idx_int = None
        if f_path and os.path.exists(f_path):
            return f_path, f_idx_int
        win_candidates: list[tuple[str, int | None]] = [
            (r"C:\\Windows\\Fonts\\simsunb.ttf", None),
            (r"C:\\Windows\\Fonts\\simsun.ttc", 1),
        ]
        for path, idx in win_candidates:
            if os.path.exists(path):
                return path, idx
        return None, None

    font_name = "Helvetica"
    font_path, font_index = resolve_font()
    if font_path:
        try:
            if font_index is not None:
                pdfmetrics.registerFont(TTFont("EmbeddedCJK", font_path, subfontIndex=font_index))
            else:
                pdfmetrics.registerFont(TTFont("EmbeddedCJK", font_path))
            font_name = "EmbeddedCJK"
        except Exception:
            font_name = "Helvetica"

    fallback_name = None
    f_path, f_index = resolve_fallback_font()
    if f_path:
        try:
            if f_index is not None:
                pdfmetrics.registerFont(TTFont("EmbeddedCJKFallback", f_path, subfontIndex=f_index))
            else:
                pdfmetrics.registerFont(TTFont("EmbeddedCJKFallback", f_path))
            fallback_name = "EmbeddedCJKFallback"
        except Exception:
            fallback_name = None

    buffer = BytesIO()
    x_margin = 20 * mm
    y_margin = 20 * mm

    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=x_margin,
        rightMargin=x_margin,
        topMargin=y_margin,
        bottomMargin=y_margin,
    )

    style = ParagraphStyle(
        name="Body",
        fontName=font_name,
        fontSize=11,
        leading=14,
        wordWrap="CJK",
    )

    def inject_fallback_runs(s: str) -> str:
        if not fallback_name:
            return htmlmod.escape(s)
        def needs_fb(code: int) -> bool:
            return (
                (0x3400 <= code <= 0x4DBF) or
                (0x4E00 <= code <= 0x9FFF) or
                (0xF900 <= code <= 0xFAFF) or
                (0x20000 <= code <= 0x2A6DF) or
                (0x2A700 <= code <= 0x2B73F) or
                (0x2B740 <= code <= 0x2B81F) or
                (0x2B820 <= code <= 0x2CEAF) or
                (0x2CEB0 <= code <= 0x2EBEF) or
                (0x30000 <= code <= 0x3134F)
            )
        out = []
        using_fb = False
        buf = []
        def flush():
            if not buf:
                return
            seg = htmlmod.escape("".join(buf))
            if using_fb:
                out.append(f"<font name=\"{fallback_name}\">" + seg + "</font>")
            else:
                out.append(seg)
            buf.clear()
        for ch in s:
            code = ord(ch)
            need_fb = needs_fb(code)
            if need_fb != using_fb:
                flush()
                using_fb = need_fb
            buf.append(ch)
        flush()
        return "".join(out)

    # 라인 원본: 타임스탬프/화자 옵션이면 SRT 사용
    lines = None
    if ts or spk:
        entries = parse_srt_entries(job_id)
        if entries:
            def make_line(e: dict) -> str:
                parts = []
                if ts and e.get("start"):
                    parts.append(f"[{e['start']}]")
                if spk and e.get("speaker"):
                    parts.append(f"[화자 {e['speaker']}]")
                prefix = (" ".join(parts) + " ") if parts else ""
                return prefix + e["text"]
            lines = [make_line(e) for e in entries]
    if lines is None:
        lines = plain.splitlines()

    story = []
    for line in lines:
        if not str(line).strip():
            story.append(Spacer(1, 6))
        else:
            safe = inject_fallback_runs(str(line)).replace("\t", "&emsp;&emsp;")
            story.append(Paragraph(safe, style))

    doc.build(story)
    buffer.seek(0)
    headers = {"Content-Disposition": f"attachment; filename=transcript_{job_id}.pdf"}
    return StreamingResponse(buffer, media_type="application/pdf", headers=headers)


# 내보내기: TXT (타임스탬프 옵션)
@app.get("/export/txt/{job_id}")
def export_txt(job_id: str, ts: int = 0, spk: int = 0):
    txt_path = os.path.join(OUTPUT_FOLDER, f"{job_id}.txt")
    if not os.path.exists(txt_path):
        raise HTTPException(status_code=404, detail="TXT 파일을 찾을 수 없습니다")
    if ts or spk:
        entries = parse_srt_entries(job_id)
        if entries:
            lines = []
            for e in entries:
                parts = []
                if ts and e.get("start"):
                    parts.append(f"[{e['start']}]")
                if spk and e.get("speaker"):
                    parts.append(f"[화자 {e['speaker']}]")
                prefix = (" ".join(parts) + " ") if parts else ""
                lines.append(prefix + e["text"])
            out = "\n".join(lines)
        else:
            with open(txt_path, "r", encoding="utf-8") as f:
                out = f.read()
    else:
        with open(txt_path, "r", encoding="utf-8") as f:
            out = f.read()
    buffer = BytesIO(out.encode("utf-8"))
    headers = {"Content-Disposition": f"attachment; filename=transcript_{job_id}.txt"}
    return StreamingResponse(buffer, media_type="text/plain; charset=utf-8", headers=headers)


# 내보내기: CSV (start,end,text)
@app.get("/export/csv/{job_id}")
def export_csv(job_id: str, spk: int = 0):
    entries = parse_srt_entries(job_id)
    if not entries:
        raise HTTPException(status_code=404, detail="SRT 파일을 찾을 수 없습니다")
    buffer = BytesIO()
    writer = csv.writer(buffer, lineterminator='\n')
    header = ["start", "end"]
    if spk:
        header.append("speaker")
    header.append("text")
    writer.writerow(header)
    for e in entries:
        row = [e["start"], e["end"]]
        if spk:
            row.append(e.get("speaker", ""))
        row.append(e["text"])
        writer.writerow(row)
    buffer.seek(0)
    headers = {"Content-Disposition": f"attachment; filename=transcript_{job_id}.csv"}
    return StreamingResponse(buffer, media_type="text/csv; charset=utf-8", headers=headers)


# 내보내기: VTT (SRT 변환)
@app.get("/export/vtt/{job_id}")
def export_vtt(job_id: str):
    srt_path = os.path.join(OUTPUT_FOLDER, f"{job_id}.srt")
    if not os.path.exists(srt_path):
        raise HTTPException(status_code=404, detail="SRT 파일을 찾을 수 없습니다")
    with open(srt_path, "r", encoding="utf-8") as f:
        srt_text = f.read()
    # 번호 라인 제거, 콤마를 점으로
    lines = []
    for line in srt_text.splitlines():
        if re.match(r"^\d+$", line.strip()):
            continue
        line = re.sub(r"(\d{2}:\d{2}:\d{2}),(\d{1,3})\s*-->\s*(\d{2}:\d{2}:\d{2}),(\d{1,3})",
                      lambda m: f"{m.group(1)}.{m.group(2).zfill(3)} --> {m.group(3)}.{m.group(4).zfill(3)}", line)
        lines.append(line)
    vtt = "WEBVTT\n\n" + "\n".join(lines)
    buffer = BytesIO(vtt.encode("utf-8"))
    headers = {"Content-Disposition": f"attachment; filename=transcript_{job_id}.vtt"}
    return StreamingResponse(buffer, media_type="text/vtt; charset=utf-8", headers=headers)